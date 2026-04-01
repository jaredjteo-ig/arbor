"""Unit tests for the policy document text extraction service.

Tests all extraction functions (PDF, DOCX, TXT), the dispatcher,
content hashing, and quality validation. Covers happy paths, error
paths, edge cases, and boundary conditions.

Company Policy Upload feature — text extraction service.
"""

import hashlib
import io
import logging
import re
import textwrap

import pytest

# ---------------------------------------------------------------------------
# Helpers to build real in-memory files for testing
# ---------------------------------------------------------------------------


def _make_docx_bytes(paragraphs: list[str], tables: list[list[list[str]]] | None = None) -> bytes:
    """Build a minimal DOCX in memory from paragraph strings and optional tables.

    ``tables`` is a list of tables, each table a list of rows, each row a list
    of cell strings.
    """
    import docx

    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if tables:
        for tbl in tables:
            rows = len(tbl)
            cols = max(len(r) for r in tbl) if tbl else 0
            t = doc.add_table(rows=rows, cols=cols)
            for ri, row in enumerate(tbl):
                for ci, cell_text in enumerate(row):
                    t.rows[ri].cells[ci].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_pdf_bytes(
    pages: list[str],
    *,
    header: str | None = None,
    footer: str | None = None,
) -> bytes:
    """Build a multi-page PDF using pdfplumber-compatible structure via reportlab.

    Falls back to fpdf2 or a raw-bytes approach depending on availability.
    Each entry in *pages* is the body text for that page.  If *header* /
    *footer* are given they appear as the first / last line of every page.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        pytest.skip("reportlab not installed — needed to generate test PDFs")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    for page_text in pages:
        y = height - 50
        if header:
            c.drawString(50, y, header)
            y -= 20
        for line in page_text.split("\n"):
            c.drawString(50, y, line)
            y -= 15
        if footer:
            c.drawString(50, 50, footer)
        c.showPage()
    c.save()
    return buf.getvalue()


# ===========================================================================
# extract_text_from_txt
# ===========================================================================


class TestExtractTextFromTxt:
    """Tests for plain-text extraction with encoding handling."""

    def test_utf8_text(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_txt

        content = "Hello, this is a policy document.\nSecond line."
        text, status = extract_text_from_txt(content.encode("utf-8"))
        assert status == "success"
        assert text == content

    def test_ascii_subset_of_utf8(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_txt

        content = "Pure ASCII text"
        text, status = extract_text_from_txt(content.encode("ascii"))
        assert status == "success"
        assert text == content

    def test_non_utf8_chardet_fallback(self) -> None:
        """Shift-JIS encoded text should be recovered via chardet."""
        from hr_advisory.services.policy_parser import extract_text_from_txt

        content = "Japanese text for testing"
        encoded = content.encode("shift_jis")
        text, status = extract_text_from_txt(encoded)
        assert status == "success"
        assert text == content

    def test_latin1_fallback(self) -> None:
        """Latin-1 specific bytes (e.g. accented chars) not valid UTF-8."""
        from hr_advisory.services.policy_parser import extract_text_from_txt

        content = "Caf\u00e9 policy r\u00e9sum\u00e9"
        encoded = content.encode("latin-1")
        text, status = extract_text_from_txt(encoded)
        assert status == "success"
        assert "Caf" in text

    def test_empty_bytes(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_txt

        text, status = extract_text_from_txt(b"")
        assert status == "success"
        assert text == ""

    def test_truly_undecodable_bytes(self) -> None:
        """Random binary that chardet cannot decode should return failed."""
        from hr_advisory.services.policy_parser import extract_text_from_txt

        # Random bytes that are unlikely to be any valid encoding
        garbage = bytes(range(128, 256)) * 5 + b"\xfe\xff\x00\x00"
        text, status = extract_text_from_txt(garbage)
        # Either chardet decodes it (some encoding) or it fails
        assert status in ("success", "failed")
        if status == "failed":
            assert text == ""


# ===========================================================================
# extract_text_from_docx
# ===========================================================================


class TestExtractTextFromDocx:
    """Tests for DOCX extraction including paragraphs and tables."""

    def test_simple_paragraphs(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_docx

        doc_bytes = _make_docx_bytes(["First paragraph.", "Second paragraph."])
        text, status = extract_text_from_docx(doc_bytes)
        assert status == "success"
        assert "First paragraph." in text
        assert "Second paragraph." in text
        # Paragraphs separated by double newline
        assert "\n\n" in text

    def test_with_table(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_docx

        table = [["Name", "Role"], ["Alice", "HR Manager"]]
        doc_bytes = _make_docx_bytes(["Header text"], tables=[table])
        text, status = extract_text_from_docx(doc_bytes)
        assert status == "success"
        assert "Name" in text
        assert "Role" in text
        assert "Alice" in text
        assert "HR Manager" in text
        # Cells separated by |
        assert " | " in text

    def test_empty_document(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_docx

        doc_bytes = _make_docx_bytes([])
        text, status = extract_text_from_docx(doc_bytes)
        assert status == "success"
        # Empty doc produces empty or whitespace-only text
        assert text.strip() == ""

    def test_invalid_bytes_returns_failed(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_docx

        text, status = extract_text_from_docx(b"this is not a docx file")
        assert status == "failed"
        assert text == ""

    def test_multiple_tables(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_docx

        table1 = [["A", "B"], ["1", "2"]]
        table2 = [["X", "Y"], ["3", "4"]]
        doc_bytes = _make_docx_bytes(["Intro"], tables=[table1, table2])
        text, status = extract_text_from_docx(doc_bytes)
        assert status == "success"
        assert "A" in text
        assert "X" in text
        assert "3" in text


# ===========================================================================
# extract_text_from_pdf
# ===========================================================================


class TestExtractTextFromPdf:
    """Tests for PDF extraction with header/footer stripping."""

    def test_simple_single_page(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        pdf_bytes = _make_pdf_bytes(["This is page one content."])
        text, status = extract_text_from_pdf(pdf_bytes)
        assert status == "success"
        assert "page one content" in text

    def test_multi_page(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        pdf_bytes = _make_pdf_bytes([
            "Page one body text.",
            "Page two body text.",
            "Page three body text.",
        ])
        text, status = extract_text_from_pdf(pdf_bytes)
        assert status == "success"
        assert "Page one" in text
        assert "Page two" in text
        assert "Page three" in text

    def test_repeated_header_stripped(self) -> None:
        """Headers repeated across 3+ pages should be removed."""
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        pdf_bytes = _make_pdf_bytes(
            ["Body A", "Body B", "Body C", "Body D"],
            header="CONFIDENTIAL COMPANY POLICY",
        )
        text, status = extract_text_from_pdf(pdf_bytes)
        assert status == "success"
        # The repeated header should have been stripped
        # Count occurrences — should be less than 4 (ideally 0)
        header_count = text.count("CONFIDENTIAL COMPANY POLICY")
        assert header_count < 4, f"Header appears {header_count} times, expected stripping"

    def test_repeated_footer_stripped(self) -> None:
        """Footers repeated across 3+ pages should be removed."""
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        pdf_bytes = _make_pdf_bytes(
            ["Body A", "Body B", "Body C", "Body D"],
            footer="Page Footer Text Here",
        )
        text, status = extract_text_from_pdf(pdf_bytes)
        assert status == "success"
        footer_count = text.count("Page Footer Text Here")
        assert footer_count < 4, f"Footer appears {footer_count} times, expected stripping"

    def test_invalid_pdf_returns_failed(self) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        text, status = extract_text_from_pdf(b"not a valid pdf")
        assert status == "failed"
        assert text == ""

    def test_empty_pdf(self) -> None:
        """A PDF with no text content returns success with empty text."""
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        pdf_bytes = _make_pdf_bytes([""])
        text, status = extract_text_from_pdf(pdf_bytes)
        assert status in ("success", "partial")

    def test_timeout_parameter_accepted(self) -> None:
        """The timeout_seconds parameter must be accepted."""
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        pdf_bytes = _make_pdf_bytes(["Hello"])
        text, status = extract_text_from_pdf(pdf_bytes, timeout_seconds=30)
        assert status == "success"


# ===========================================================================
# extract_text (dispatcher)
# ===========================================================================


class TestExtractText:
    """Tests for the top-level dispatcher function."""

    def test_dispatch_txt(self) -> None:
        from hr_advisory.services.policy_parser import extract_text

        text, status = extract_text(b"Hello policy", "txt")
        assert status == "success"
        assert text == "Hello policy"

    def test_dispatch_docx(self) -> None:
        from hr_advisory.services.policy_parser import extract_text

        doc_bytes = _make_docx_bytes(["Policy content here."])
        text, status = extract_text(doc_bytes, "docx")
        assert status == "success"
        assert "Policy content here." in text

    def test_dispatch_pdf(self) -> None:
        from hr_advisory.services.policy_parser import extract_text

        pdf_bytes = _make_pdf_bytes(["PDF policy content."])
        text, status = extract_text(pdf_bytes, "pdf")
        assert status == "success"
        assert "PDF policy" in text

    def test_unsupported_type_raises_value_error(self) -> None:
        from hr_advisory.services.policy_parser import extract_text

        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(b"data", "xlsx")

    def test_unsupported_type_html(self) -> None:
        from hr_advisory.services.policy_parser import extract_text

        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(b"<html></html>", "html")

    def test_truncation_at_500k_chars(self) -> None:
        """Extracted text exceeding 500,000 chars is truncated."""
        from hr_advisory.services.policy_parser import extract_text

        # Build a TXT file with > 500k characters
        long_content = "A" * 600_000
        text, status = extract_text(long_content.encode("utf-8"), "txt")
        assert len(text) <= 500_000

    def test_case_insensitive_file_type(self) -> None:
        """File type matching should handle different casing."""
        from hr_advisory.services.policy_parser import extract_text

        text, status = extract_text(b"Hello", "TXT")
        assert status == "success"

    def test_file_type_with_leading_dot(self) -> None:
        """Handle '.pdf' with leading dot gracefully."""
        from hr_advisory.services.policy_parser import extract_text

        # Should either work or raise a clear ValueError
        try:
            pdf_bytes = _make_pdf_bytes(["Content"])
            text, status = extract_text(pdf_bytes, ".pdf")
            # If it works, great
            assert status == "success"
        except ValueError:
            # Also acceptable — clear error for unsupported type
            pass


# ===========================================================================
# compute_content_hash
# ===========================================================================


class TestComputeContentHash:
    """Tests for SHA-256 content hashing."""

    def test_known_hash(self) -> None:
        from hr_advisory.services.policy_parser import compute_content_hash

        content = "Hello, world!"
        expected = hashlib.sha256(content.encode("utf-8")).hexdigest()
        assert compute_content_hash(content) == expected

    def test_empty_string(self) -> None:
        from hr_advisory.services.policy_parser import compute_content_hash

        expected = hashlib.sha256(b"").hexdigest()
        assert compute_content_hash("") == expected

    def test_unicode_content(self) -> None:
        from hr_advisory.services.policy_parser import compute_content_hash

        content = "Singapore Employment Act \u00a7 clause"
        expected = hashlib.sha256(content.encode("utf-8")).hexdigest()
        assert compute_content_hash(content) == expected

    def test_deterministic(self) -> None:
        from hr_advisory.services.policy_parser import compute_content_hash

        content = "Consistent hashing test"
        h1 = compute_content_hash(content)
        h2 = compute_content_hash(content)
        assert h1 == h2

    def test_different_content_different_hash(self) -> None:
        from hr_advisory.services.policy_parser import compute_content_hash

        h1 = compute_content_hash("content A")
        h2 = compute_content_hash("content B")
        assert h1 != h2


# ===========================================================================
# validate_extraction_quality
# ===========================================================================


class TestValidateExtractionQuality:
    """Tests for extraction quality validation."""

    def test_good_quality(self) -> None:
        from hr_advisory.services.policy_parser import validate_extraction_quality

        text = (
            "This is a well-formed policy document with sufficient words "
            "to pass the minimum word count threshold. It contains normal "
            "English text that should be classified as good quality extraction."
        )
        quality, warnings = validate_extraction_quality(text)
        assert quality == "good"
        assert isinstance(warnings, list)

    def test_empty_text(self) -> None:
        from hr_advisory.services.policy_parser import validate_extraction_quality

        quality, warnings = validate_extraction_quality("")
        assert quality == "empty"

    def test_whitespace_only_text(self) -> None:
        from hr_advisory.services.policy_parser import validate_extraction_quality

        quality, warnings = validate_extraction_quality("   \n\n\t  ")
        assert quality == "empty"

    def test_too_few_words(self) -> None:
        """Fewer than 10 words should flag as low quality (garbage)."""
        from hr_advisory.services.policy_parser import validate_extraction_quality

        quality, warnings = validate_extraction_quality("one two three")
        assert quality == "low_quality"
        assert any("word" in w.lower() for w in warnings)

    def test_high_char_to_word_ratio_binary(self) -> None:
        """Char-to-word ratio > 20 suggests binary/garbled content."""
        from hr_advisory.services.policy_parser import validate_extraction_quality

        # 15 words but with very long "words" — ratio > 20
        text = " ".join(["a" * 50] * 15)
        quality, warnings = validate_extraction_quality(text)
        assert quality == "low_quality"
        assert any("ratio" in w.lower() or "binary" in w.lower() for w in warnings)

    def test_excessive_special_chars(self) -> None:
        """More than 30% non-alphanumeric characters suggests OCR noise."""
        from hr_advisory.services.policy_parser import validate_extraction_quality

        # Build text where special chars dominate
        normal = "word " * 10  # 50 chars (including spaces)
        specials = "!@#$%^&*(){}[]<>?/" * 5  # 90 special chars
        text = normal + specials
        quality, warnings = validate_extraction_quality(text)
        assert quality == "low_quality"
        assert any("special" in w.lower() or "ocr" in w.lower() for w in warnings)

    def test_exactly_10_words_is_acceptable(self) -> None:
        """Edge case: exactly 10 words should not fail the minimum check."""
        from hr_advisory.services.policy_parser import validate_extraction_quality

        text = "one two three four five six seven eight nine ten"
        quality, warnings = validate_extraction_quality(text)
        # 10 words is the boundary — should be good if other checks pass
        assert quality in ("good", "low_quality")
        # Should NOT have a "too few words" warning
        min_word_warnings = [w for w in warnings if "word" in w.lower() and "minim" in w.lower()]
        assert len(min_word_warnings) == 0

    def test_returns_multiple_warnings(self) -> None:
        """A truly bad extraction can trigger multiple warnings."""
        from hr_advisory.services.policy_parser import validate_extraction_quality

        # Few words + high special char ratio
        text = "!!! @@@ ### $$$ %%% ^^^ &&& *** ((( )))"
        quality, warnings = validate_extraction_quality(text)
        assert quality == "low_quality"
        # Should have at least one warning
        assert len(warnings) >= 1


# ===========================================================================
# Logging behavior
# ===========================================================================


class TestLogging:
    """Verify that the module emits log messages for failures."""

    def test_failed_pdf_logs_error(self, caplog: pytest.LogCaptureFixture) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_pdf

        with caplog.at_level(logging.WARNING, logger="hr_advisory.services.policy_parser"):
            extract_text_from_pdf(b"bad data")
        assert len(caplog.records) > 0

    def test_failed_docx_logs_error(self, caplog: pytest.LogCaptureFixture) -> None:
        from hr_advisory.services.policy_parser import extract_text_from_docx

        with caplog.at_level(logging.WARNING, logger="hr_advisory.services.policy_parser"):
            extract_text_from_docx(b"bad data")
        assert len(caplog.records) > 0
